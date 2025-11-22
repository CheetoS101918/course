import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import sys
import numpy as np


def simple_clean_zeros(df):
    """Превращает целые float-ы (5.0, 10.0 и т.д.) в int, остальное оставляет как есть"""
    def format_number(x):
        if pd.isna(x):
            return x
        try:
            num = float(x)
            # Если число целое — возвращаем int, иначе оставляем float
            return int(num) if num.is_integer() else num
        except (ValueError, TypeError, OverflowError):
            return x
    return df.map(format_number)  # map, а не applymap (applymap устарел в новых pandas)
    

def value_to_string(val):
    """Надёжно убирает .0 у целых чисел, даже если они пришли как float"""
    if pd.isna(val):
        return ""
    if isinstance(val, float):
        if val.is_integer():
            return str(int(val))
        return str(val).rstrip('0').rstrip('.')  # убирает лишние нули и точку у 10.50 → 10.5
    return str(val)


def create_custom_style(doc):
    """Создает изолированный стиль для таблицы без отступов"""
    styles = doc.styles
    # Создаем новый стиль абзаца
    style_name = 'CsvTableText'

    try:
        custom_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        custom_style.base_style = styles['Normal']
    except:
        custom_style = styles[style_name]  # Если вдруг уже есть

    # Настраиваем стиль жестко
    p_fmt = custom_style.paragraph_format
    p_fmt.first_line_indent = Pt(0)
    p_fmt.left_indent = Pt(0)
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)

    # Можно добавить шрифт, чтобы он тоже не слетал
    font = custom_style.font
    font.name = 'Times New Roman'  # Или какой вам нужен
    font.size = Pt(14)

    return style_name


def create_docx_from_csv(csv_file, output_file=None):
    """
    Создаёт DOCX-документ с таблицей из CSV-файла.

    Args:
        csv_file (str): Путь к CSV-файлу (с разделителем ';').
        output_file (str, optional): Путь к выходному DOCX. Если None,
                                     генерирует имя на основе CSV.

    Returns:
        str: Путь к сохранённому файлу.
    """
    # Шаг 1: Читаем CSV
    df1 = pd.read_csv(csv_file, sep=';')

    df = simple_clean_zeros(df1) # очистка от грёбаных .0

    # Шаг 2: Генерируем имя выходного файла, если не указано
    if output_file is None:
        output_file = csv_file.replace('.csv', '_table.docx')

    # Шаг 3: Создаём DOCX
    doc = Document()

    table_style_name = create_custom_style(doc)

    # Шаг 4: Добавляем таблицу
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'  # Рамки

    # Шаг 5: Заполняем заголовки
    for j, column in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(column)
        # !!! 2. Применяем стиль
        cell.paragraphs[0].style = doc.styles[table_style_name]
        cell.paragraphs[0].runs[0].bold = True

    # Шаг 6: Заполняем данные (NaN -> пустая ячейка)
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = value_to_string(value)      # ← вот тут магия
            cell.paragraphs[0].style = doc.styles[table_style_name]

    # Шаг 7: Настраиваем ширину столбцов
    table.columns[0].width = Inches(2.5)  # Шире для названий
    for j in range(1, len(table.columns)):
        table.columns[j].width = Inches(1.2)

    # Шаг 8: Сохраняем
    doc.save(output_file)
    print(f"Таблица сохранена в {output_file}")
    return output_file


if __name__ == "__main__":
    # Парсим аргументы из командной строки
    if len(sys.argv) != 2:
        print("Использование: python csv_to_docx.py <путь_к_csv_файлу>")
        print("Пример: python csv_to_docx.py sebestoimost_structure.csv")
        sys.exit(1)

    csv_file = sys.argv[1]
    create_docx_from_csv(csv_file)